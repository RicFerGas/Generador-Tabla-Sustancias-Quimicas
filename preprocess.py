#preprocess_hds.py
import os
import sys
import PyPDF2
import spacy
import fitz
import pytesseract
from docx import Document
from pdf2image import convert_from_path

class DocumentPreprocessor:
    """
    A class for preprocessing and extracting text from PDF and Word documents.
    
    This class provides functionality to:
    - Validate document formats (.pdf, .doc, .docx)
    - Extract text using multiple methods (native PDF, OCR, Word)
    - Validate extracted text quality
    - Process documents with Spanish language support
    
    Attributes:
        supported_extensions (list): List of supported file extensions ['.pdf', '.docx', '.doc']
        nlp: Spacy language model for text processing
        file_path (str): Path to the current document being processed
    """
    
    def __init__(self):
            # Initialize spacy model
            if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(os.path.abspath(__file__))
            self.supported_extensions = ['.pdf', '.docx', '.doc']
            model_relative_path = os.path.join('models',
                                            'xx_ent_wiki_sm',
                                            'xx_ent_wiki_sm',
                                            'xx_ent_wiki_sm-3.8.0')
            self.file_path = None
            model_path = os.path.join(base_path, model_relative_path)
            self.nlp = spacy.load(model_path)
    def is_supported_file(self) -> bool:
        """
        Checks if the file is a supported document type.

        Args:
            file_path (str): The path to the file.

        Returns:
            bool: True if the file is supported, False otherwise.
        """
        return any(self.file_path.lower().endswith(ext) for ext in self.supported_extensions)

    def is_valid_pdf(self) -> bool:
        """
        Validates PDF file structure and content.
        
        Checks:
        - Valid PDF header
        - File can be opened
        - Contains at least one page
        - Has extractable text
        
        Returns:
            bool: True if PDF is valid and contains extractable content
        """
        try:
            # Check if it starts with %PDF-
            with open(self.file_path, 'rb') as file:
                header = file.read(4)
                if header != b'%PDF':
                    print(f"File is not a valid PDF: {self.file_path}")
                    return False

                # Attempt to open and read the PDF
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # Check if there are any pages
                if num_pages == 0:
                    print(f"PDF has no pages: {self.file_path}")
                    return False
                
                
                print(f"PDF is valid: {self.file_path} (pages: {num_pages})")
                return True
                
        except Exception as e:
            print(f"Invalid PDF {self.file_path}: {e}")
            return False
    def validate_extracted_text(self,text: str, threshold: float = 0.6) -> bool:
        """
        Validates the extracted text from a PDF based on the word count threshold.

        Args:
            text (str): The extracted text from the PDF.
            threshold (float): The minimum word count rate to consider the text valid.

        Returns:
            bool: True if the text is valid, False otherwise.
        """
        try:
            doc = self.nlp(text)
            word_count = sum(1 for token in doc if token.is_alpha)
            total_count = len(doc)
            print(f"word count: {word_count}, total count: {total_count}, rate: {word_count / total_count}")
            if total_count == 0:
                return False
            return (word_count / total_count) >= threshold
        except Exception as e:
            print(f"Error in text validation: {e}")
            return False
    def extract_text_from_word(self) -> str:
        """
        Extracts text from a Word document using the docx library.

        Args:
            docx_path (str): The path to the Word document.

        Returns:
            str: The extracted text from the Word document.
        """
        try:
            
            doc = Document(self.file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        content_parts.append(' | '.join(row_text))

            # Join all content with newlines
            text = '\n'.join(content_parts)
            
            # Basic cleaning
            text = ' '.join(text.split())  # Remove extra whitespace
            return text if text else None
        except Exception as e:
            print(f"Error extracting text from {self.file_path}: {str(e)}")
            return None
    def extract_text_from_pdf_native(self) -> str:
        """
        Extracts text from PDF using PyMuPDF.
        """
        try:
            text = ""
            with fitz.open(self.file_path) as pdf:
                for page_num in range(len(pdf)):
                    page = pdf.load_page(page_num)
                    text += page.get_text()
            return text
        except Exception as e:
            print(f"Error extracting text using PyMuPDF: {e}")
            return ""

    def extract_text_using_ocr(self) -> str:
        """
        Extracts text from PDF using OCR.
        """
        try:
            images = convert_from_path(self.file_path)
            extracted_text = ""
            for image in images:
                extracted_text += pytesseract.image_to_string(image, lang='spa')
            return extracted_text
        except Exception as e:
            print(f"Error during OCR: {e}")
            return ""

    def extract_text(self,document_path:str) -> str:
        """
        Main method to extract text from any supported document.
        """
        self.file_path=document_path
        if not self.is_supported_file():
            return "Unsupported file format"

        if self.file_path.lower().endswith('.pdf'):
            if not self.is_valid_pdf():
                return "Invalid PDF file"
            
            text = self.extract_text_from_pdf_native()
            if text and self.validate_extracted_text(text):
                return text
                
            print("Native PDF extraction failed, trying OCR...")
            text = self.extract_text_using_ocr()
            return text if text else "Text extraction failed"
            
        elif self.file_path.lower().endswith(('.docx', '.doc')):
            return self.extract_text_from_word() or "Text extraction failed"        
if __name__ == "__main__":
    file_path = "ejemplo/mini_test/65-FO-SP80-40 Wear Guard 2000.docx"
    doc_preprocessor = DocumentPreprocessor()
    extracted_text = doc_preprocessor.extract_text(file_path)
    print(extracted_text)